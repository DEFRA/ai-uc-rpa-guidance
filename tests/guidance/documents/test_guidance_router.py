"""Tests for the guidance API endpoints."""

import io
import json
import uuid
from unittest.mock import AsyncMock

import botocore.exceptions
import docx
import fastapi.testclient
import pytest

import app.entrypoints.fastapi
from app.guidance.documents import api_schemas, parser, s3_repository
from app.guidance.documents import dependencies as guidance_dependencies
from app.guidance.documents import models as guidance_models


def _no_such_key_error() -> botocore.exceptions.ClientError:
    return botocore.exceptions.ClientError(
        error_response={
            "Error": {"Code": "NoSuchKey", "Message": "The key does not exist"}
        },
        operation_name="GetObject",
    )


@pytest.fixture
def mock_guidance_service() -> AsyncMock:
    """Create a mock guidance service."""
    return AsyncMock()


@pytest.fixture
def mock_s3_repo() -> AsyncMock:
    """Create a mock S3 repository."""
    return AsyncMock(spec=s3_repository.AbstractGuidanceStorageRepository)


@pytest.fixture
def client_with_mocks(
    mock_guidance_service: AsyncMock,
) -> fastapi.testclient.TestClient:
    """Create a test client with mocked guidance service."""
    test_app = app.entrypoints.fastapi.app
    test_app.dependency_overrides[guidance_dependencies.get_guidance_service] = lambda: (
        mock_guidance_service
    )
    return fastapi.testclient.TestClient(test_app)


@pytest.fixture
def client_with_s3(
    mock_s3_repo: AsyncMock,
) -> fastapi.testclient.TestClient:
    """Create a test client with a mocked S3 repository."""
    test_app = app.entrypoints.fastapi.app
    test_app.dependency_overrides[guidance_dependencies.get_s3_repository] = lambda: (
        mock_s3_repo
    )
    return fastapi.testclient.TestClient(test_app)


class TestInitiateEndpoint:
    """Test the POST /guidance/documents endpoint."""

    def test_initiate_success(
        self,
        client_with_mocks: fastapi.testclient.TestClient,
        mock_guidance_service: AsyncMock,
    ) -> None:
        """Test initiating a document upload successfully."""
        upload_id = "507f1f77bcf86cd799439011"
        mock_guidance_service.initiate_upload.return_value = upload_id

        response = client_with_mocks.post(
            "/guidance/documents",
            json={
                "title": "Test Document",
                "description": "A test document upload",
                "redirect": "http://example.com/return",
            },
        )

        assert response.status_code == 201
        data = response.json()
        assert "uploadId" in data
        assert data["uploadId"] == upload_id

    def test_initiate_without_title(
        self,
        client_with_mocks: fastapi.testclient.TestClient,
        mock_guidance_service: AsyncMock,
    ) -> None:
        """Title is optional — inferred from document during parsing."""
        upload_id = "507f191e810c19729de860ea"
        mock_guidance_service.initiate_upload.return_value = upload_id

        response = client_with_mocks.post(
            "/guidance/documents",
            json={
                "description": "No title provided",
                "redirect": "http://example.com/return",
            },
        )

        assert response.status_code == 201

    def test_initiate_with_missing_redirect(
        self, client_with_mocks: fastapi.testclient.TestClient
    ) -> None:
        """Test initiate without required redirect field."""
        response = client_with_mocks.post(
            "/guidance/documents",
            json={
                "title": "Test",
                "description": "Missing redirect",
            },
        )

        assert response.status_code == 422

    def test_initiate_with_service_error(
        self,
        client_with_mocks: fastapi.testclient.TestClient,
        mock_guidance_service: AsyncMock,
    ) -> None:
        """Test initiate when service raises an exception."""
        mock_guidance_service.initiate_upload.side_effect = Exception(
            "CDP uploader error"
        )

        response = client_with_mocks.post(
            "/guidance/documents",
            json={
                "title": "Test Document",
                "description": "A test document upload",
                "redirect": "http://example.com/return",
            },
        )

        assert response.status_code == 502


class TestCallbackEndpoint:
    """Test the POST /guidance/documents/{document_id}/callback endpoint."""

    def test_callback_with_valid_payload(
        self,
        client_with_mocks: fastapi.testclient.TestClient,
        mock_guidance_service: AsyncMock,
    ) -> None:
        """Test callback handler with valid CDP uploader payload."""
        mock_guidance_service.handle_callback.return_value = None

        document_id = "12345678-1234-5678-1234-567812345678"
        response = client_with_mocks.post(
            f"/guidance/documents/{document_id}/callback",
            json={
                "uploadStatus": "completed",
                "form": {
                    "file1": {
                        "fileId": "file-1",
                        "filename": "test.pdf",
                        "fileStatus": "completed",
                        "contentLength": 1024,
                        "checksumSha256": "abc123",
                        "s3Key": "test.pdf",
                        "s3Bucket": "guidance-bucket",
                    }
                },
            },
        )

        assert response.status_code == 204

    def test_callback_with_invalid_document_id(
        self, client_with_mocks: fastapi.testclient.TestClient
    ) -> None:
        """Test callback with a non-UUID document ID is rejected at the router level."""
        response = client_with_mocks.post(
            "/guidance/documents/not-a-uuid/callback",
            json={
                "uploadStatus": "completed",
                "form": {},
            },
        )

        assert response.status_code == 422

    def test_callback_with_nonexistent_document(
        self,
        client_with_mocks: fastapi.testclient.TestClient,
        mock_guidance_service: AsyncMock,
    ) -> None:
        """Test callback with document ID that doesn't exist."""
        mock_guidance_service.handle_callback.side_effect = ValueError(
            "Document not found"
        )

        document_id = "12345678-1234-5678-1234-567812345678"
        response = client_with_mocks.post(
            f"/guidance/documents/{document_id}/callback",
            json={
                "uploadStatus": "completed",
                "form": {},
            },
        )

        assert response.status_code == 404


class TestListDocumentsEndpoint:
    """Test the GET /guidance/documents endpoint."""

    def test_list_documents_success(
        self,
        client_with_mocks: fastapi.testclient.TestClient,
        mock_guidance_service: AsyncMock,
    ) -> None:
        """Test listing documents with pagination."""
        mock_guidance_service.list_documents.return_value = (
            api_schemas.DocumentListResponse(
                items=[],
                total=0,
                page=1,
                page_size=10,
            )
        )

        response = client_with_mocks.get("/guidance/documents?page=1&page_size=10")

        assert response.status_code == 200
        data = response.json()
        assert "items" in data
        assert "total" in data
        assert "page" in data
        assert "pageSize" in data
        assert data["page"] == 1
        assert data["pageSize"] == 10

    def test_list_documents_with_invalid_page(
        self, client_with_mocks: fastapi.testclient.TestClient
    ) -> None:
        """Test that invalid page number is rejected."""
        response = client_with_mocks.get("/guidance/documents?page=0&page_size=10")

        assert response.status_code == 422

    def test_list_documents_with_invalid_page_size(
        self, client_with_mocks: fastapi.testclient.TestClient
    ) -> None:
        """Test that invalid page size is rejected."""
        response = client_with_mocks.get("/guidance/documents?page=1&page_size=1000")

        assert response.status_code == 422

    def test_list_documents_with_defaults(
        self,
        client_with_mocks: fastapi.testclient.TestClient,
        mock_guidance_service: AsyncMock,
    ) -> None:
        """Test listing documents with default pagination parameters."""
        mock_guidance_service.list_documents.return_value = (
            api_schemas.DocumentListResponse(
                items=[],
                total=0,
                page=1,
                page_size=10,
            )
        )

        response = client_with_mocks.get("/guidance/documents")

        assert response.status_code == 200


class TestManifestEndpoint:
    """Test GET /guidance/documents/{document_id}/manifest."""

    _DOCUMENT_ID = "12345678-1234-5678-1234-567812345678"
    _MANIFEST_JSON = (
        '{"document_id": "12345678-1234-5678-1234-567812345678", "title": "Test", "sections": ['
        '{"number": "1", "heading": "Intro", "level": 1, "parent": null, "children": [], "links": []}'
        "]}"
    )

    def test_returns_manifest(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        mock_s3_repo.download_manifest.return_value = self._MANIFEST_JSON

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/manifest"
        )

        assert response.status_code == 200
        data = response.json()
        assert data["documentId"] == self._DOCUMENT_ID
        assert len(data["sections"]) == 1
        assert data["sections"][0]["number"] == "1"

    def test_returns_404_when_manifest_missing(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        mock_s3_repo.download_manifest.side_effect = _no_such_key_error()

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/manifest"
        )

        assert response.status_code == 404

    def test_returns_422_for_invalid_document_id(
        self,
        client_with_s3: fastapi.testclient.TestClient,
    ) -> None:
        response = client_with_s3.get("/guidance/documents/not-a-uuid/manifest")

        assert response.status_code == 422


class TestContentEndpoint:
    """Test GET /guidance/documents/{document_id}/content."""

    _DOCUMENT_ID = "12345678-1234-5678-1234-567812345678"

    def test_returns_full_document_markdown(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        mock_s3_repo.download_content.return_value = (
            "## 1 Intro\n\nContent.\n\n## 2 Next\n\nMore content."
        )

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/content"
        )

        assert response.status_code == 200
        assert "text/markdown" in response.headers["content-type"]
        assert "## 1 Intro" in response.text
        assert "## 2 Next" in response.text

    def test_returns_404_when_content_missing(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        mock_s3_repo.download_content.side_effect = _no_such_key_error()

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/content"
        )

        assert response.status_code == 404

    def test_returns_422_for_invalid_document_id(
        self,
        client_with_s3: fastapi.testclient.TestClient,
    ) -> None:
        response = client_with_s3.get("/guidance/documents/not-a-uuid/content")

        assert response.status_code == 422


class TestContentEndpointServesParsedTitle:
    """The /content Markdown must be headed by the title as it appears on the page.

    Drives a real .docx through the parse pipeline and serves what the pipeline
    stored, so the endpoint is asserted against genuinely parsed output rather
    than a hand-written fixture.
    """

    _DOCUMENT_ID = "12345678-1234-5678-1234-567812345678"

    _SCHEME_LINE = "Sustainable Farming Incentive 2023 (SFI23)"
    _SUBJECT_LINES = (
        "Parcel ID not linked to Single Business Identifier SBI)",
        "In",
        "SITI Tenure Guidance",
    )
    _EXPECTED_TITLE = (
        "Sustainable Farming Incentive 2023 (SFI23) — "
        "Parcel ID not linked to Single Business Identifier SBI) "
        "In SITI Tenure Guidance"
    )

    @staticmethod
    def _cover_page_docx() -> bytes:
        """A .docx shaped like the RPA guidance template: Title block, then body.

        The scheme name, a blank line, then the subject hand-wrapped over three
        Title paragraphs — the layout of the real guidance documents.
        """
        doc = docx.Document()
        doc.core_properties.title = "Design Team Guidance Template"
        cls = TestContentEndpointServesParsedTitle
        doc.add_paragraph(cls._SCHEME_LINE).style = doc.styles["Title"]
        doc.add_paragraph("")
        for line in cls._SUBJECT_LINES:
            doc.add_paragraph(line).style = doc.styles["Title"]
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("This guide explains the process.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def _parsed_content(self) -> str:
        """Run the pipeline over the .docx and return the Markdown it stored."""
        s3_repo = AsyncMock(spec=s3_repository.AbstractGuidanceStorageRepository)
        s3_repo.download_docx.return_value = self._cover_page_docx()

        document = guidance_models.GuidanceDocument(
            id=uuid.UUID(self._DOCUMENT_ID),
            title=None,
            status=guidance_models.ExtractionStatus.PROCESSING,
            path="s3://source-bucket/doc-id/file-id",
        )
        result = await parser.PipelineDocumentParser(s3_repo).parse(document)

        assert result.status == guidance_models.ExtractionStatus.COMPLETE
        return str(s3_repo.upload_content.await_args.args[1])

    async def test_content_is_headed_by_the_full_cover_title(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        mock_s3_repo.download_content.return_value = await self._parsed_content()

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/content"
        )

        assert response.status_code == 200
        assert response.text.splitlines()[0] == f"# {self._EXPECTED_TITLE}"

    async def test_content_title_is_not_truncated_to_its_first_line(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        """Guards the specific regression: only the first Title paragraph served."""
        mock_s3_repo.download_content.return_value = await self._parsed_content()

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/content"
        )

        heading = response.text.splitlines()[0]
        assert heading != "# Sustainable Farming Incentive 2023 (SFI23)"
        assert "SITI Tenure Guidance" in heading


class TestSectionEndpoint:
    """Test GET /guidance/documents/{document_id}/sections/{section_number}."""

    _DOCUMENT_ID = "12345678-1234-5678-1234-567812345678"

    def test_returns_section_markdown(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        mock_s3_repo.download_section.return_value = "## 1 Intro\n\nContent here."

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/1"
        )

        assert response.status_code == 200
        assert "text/markdown" in response.headers["content-type"]
        assert "## 1 Intro" in response.text

    def test_returns_section_with_dotted_number(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        mock_s3_repo.download_section.return_value = "### 1.2.3 Deep\n\nContent."

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/1.2.3"
        )

        assert response.status_code == 200

    def test_returns_404_when_section_missing(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        mock_s3_repo.download_section.side_effect = _no_such_key_error()

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/99"
        )

        assert response.status_code == 404

    def test_returns_422_for_invalid_section_number(
        self,
        client_with_s3: fastapi.testclient.TestClient,
    ) -> None:
        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/not-a-number"
        )

        assert response.status_code == 422

    def test_returns_422_for_section_with_slash(
        self,
        client_with_s3: fastapi.testclient.TestClient,
    ) -> None:
        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/1/2"
        )

        assert response.status_code in {404, 422}


class TestSectionEndpointChildren:
    """Test GET /sections/{section_number}?children=true.

    Acceptance criterion: the response must match slicing the section (plus its
    descendants) directly out of the full /content document — not merely match
    whatever the implementation happens to concatenate.
    """

    _DOCUMENT_ID = "12345678-1234-5678-1234-567812345678"

    # Chunks as they are individually stored (no manifest boilerplate, no
    # trailing newline) — mirrors how download_section results look in S3.
    _SECTION_1 = "## 1 Intro\n\nA."
    _SECTION_2 = "## 2 Middle\n\nB."
    _SECTION_2_1 = "### 2.1 Sub One\n\nC."
    _SECTION_2_2 = "### 2.2 Sub Two\n\nD."
    _SECTION_3 = "## 3 End\n\nE."

    _MANIFEST_JSON = (
        '{"document_id": "12345678-1234-5678-1234-567812345678", "title": "Test Document", '
        '"sections": ['
        '{"number": "1", "heading": "Intro", "level": 1, "parent": null, "children": [], "links": []},'
        '{"number": "2", "heading": "Middle", "level": 1, "parent": null, "children": ["2.1", "2.2"], "links": []},'
        '{"number": "2.1", "heading": "Sub One", "level": 2, "parent": "2", "children": [], "links": []},'
        '{"number": "2.2", "heading": "Sub Two", "level": 2, "parent": "2", "children": [], "links": []},'
        '{"number": "3", "heading": "End", "level": 1, "parent": null, "children": [], "links": []}'
        "]}"
    )

    @classmethod
    def _full_content(cls) -> str:
        """Build the full /content document the same way it is independently verified to be built."""
        body = "\n\n".join(
            [
                cls._SECTION_1,
                cls._SECTION_2,
                cls._SECTION_2_1,
                cls._SECTION_2_2,
                cls._SECTION_3,
            ]
        )
        return f"# Test Document\n\n{body}\n"

    @staticmethod
    def _section_lookup(document_id: object, section_number: str) -> str:  # noqa: ARG004
        lookup = {
            "1": TestSectionEndpointChildren._SECTION_1,
            "2": TestSectionEndpointChildren._SECTION_2,
            "2.1": TestSectionEndpointChildren._SECTION_2_1,
            "2.2": TestSectionEndpointChildren._SECTION_2_2,
            "3": TestSectionEndpointChildren._SECTION_3,
        }
        return lookup[section_number]

    def _mock_repo(self, mock_s3_repo: AsyncMock) -> None:
        mock_s3_repo.download_manifest.return_value = self._MANIFEST_JSON
        mock_s3_repo.download_content.return_value = self._full_content()
        mock_s3_repo.download_section.side_effect = self._section_lookup

    def test_children_matches_slice_of_full_content(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        self._mock_repo(mock_s3_repo)

        full_content = self._full_content()
        start = full_content.index("## 2 Middle")
        end = full_content.index("\n\n## 3 End")
        expected = full_content[start:end] + "\n"

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/2?children=true"
        )

        assert response.status_code == 200
        assert response.text == expected

    def test_children_at_end_of_document_matches_slice_to_end(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        self._mock_repo(mock_s3_repo)

        full_content = self._full_content()
        start = full_content.index("## 3 End")
        expected = full_content[start:]

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/3?children=true"
        )

        assert response.status_code == 200
        assert response.text == expected

    def test_children_false_is_the_default(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        self._mock_repo(mock_s3_repo)

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/2"
        )

        assert response.status_code == 200
        assert response.text == self._SECTION_2
        mock_s3_repo.download_manifest.assert_not_called()

    def test_children_false_explicit(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        self._mock_repo(mock_s3_repo)

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/2?children=false"
        )

        assert response.status_code == 200
        assert response.text == self._SECTION_2

    def test_children_with_no_children_matches_plain_section(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        self._mock_repo(mock_s3_repo)

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/1?children=true"
        )

        assert response.status_code == 200
        assert response.text == self._SECTION_1 + "\n"

    def test_children_returns_404_for_unknown_section(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        self._mock_repo(mock_s3_repo)
        mock_s3_repo.download_section.side_effect = _no_such_key_error()

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/sections/99?children=true"
        )

        assert response.status_code == 404


class TestImageEndpoint:
    """Test GET /guidance/documents/{document_id}/images/{filename}."""

    _DOCUMENT_ID = "12345678-1234-5678-1234-567812345678"

    def test_returns_image_bytes(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        image_data = b"\x89PNG\r\n\x1a\n"
        mock_s3_repo.download_image.return_value = image_data

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/images/img_1.png"
        )

        assert response.status_code == 200
        assert response.content == image_data
        assert response.headers["content-type"] == "image/png"

    def test_returns_correct_content_type_for_jpeg(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        mock_s3_repo.download_image.return_value = b"jpeg_data"

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/images/img_2.jpeg"
        )

        assert response.status_code == 200
        assert response.headers["content-type"] == "image/jpeg"

    def test_returns_404_when_image_missing(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
    ) -> None:
        mock_s3_repo.download_image.side_effect = _no_such_key_error()

        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/images/img_1.png"
        )

        assert response.status_code == 404

    def test_returns_422_for_filename_starting_with_dash(
        self,
        client_with_s3: fastapi.testclient.TestClient,
    ) -> None:
        response = client_with_s3.get(
            f"/guidance/documents/{self._DOCUMENT_ID}/images/-secret.png"
        )

        assert response.status_code == 422

    def test_returns_422_for_invalid_document_id(
        self,
        client_with_s3: fastapi.testclient.TestClient,
    ) -> None:
        response = client_with_s3.get("/guidance/documents/not-a-uuid/images/img_1.png")

        assert response.status_code == 422


class TestUpdateSectionEndpoint:
    """PUT /guidance/documents/{id}/sections/{number} applies an editor's correction."""

    MANIFEST = {
        "document_id": "aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee",
        "title": "SFI Parcel Guidance",
        "sections": [
            {
                "number": "1",
                "heading": "Overview",
                "level": 1,
                "parent": None,
                "children": [],
                "links": [],
            },
            {
                "number": "7",
                "heading": "Email template",
                "level": 1,
                "parent": None,
                "children": [],
                "links": [],
            },
        ],
    }

    @pytest.fixture
    def document_id(self) -> str:
        return self.MANIFEST["document_id"]

    @pytest.fixture(autouse=True)
    def _prime_storage(self, mock_s3_repo: AsyncMock) -> None:
        """Prime the manifest and section files a write path reads back."""
        mock_s3_repo.download_manifest.return_value = json.dumps(self.MANIFEST)
        stored = {
            "1": "## 1 Overview\n\nOriginal text.\n",
            "7": "## 7 Email template\n\nSVBI is wrong.\n",
        }

        async def download_section(_document_id: uuid.UUID, number: str) -> str:
            return stored[number]

        async def upload_section(
            _document_id: uuid.UUID, number: str, markdown: str
        ) -> None:
            stored[number] = markdown

        mock_s3_repo.download_section.side_effect = download_section
        mock_s3_repo.upload_section.side_effect = upload_section

    def test_returns_204_and_writes_the_section(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
        document_id: str,
    ) -> None:
        response = client_with_s3.put(
            f"/guidance/documents/{document_id}/sections/7",
            json={"heading": "Email template", "markdown": "SBI is correct."},
        )

        assert response.status_code == 204
        mock_s3_repo.upload_section.assert_awaited_once_with(
            uuid.UUID(document_id),
            "7",
            "## 7 Email template\n\nSBI is correct.\n",
        )

    def test_regenerates_content_document(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
        document_id: str,
    ) -> None:
        """The review checker reads content.md, so it must not go stale."""
        client_with_s3.put(
            f"/guidance/documents/{document_id}/sections/7",
            json={"heading": "Email template", "markdown": "SBI is correct."},
        )

        content = str(mock_s3_repo.upload_content.await_args.args[1])
        assert "SBI is correct." in content
        assert "SVBI" not in content

    def test_accepts_snake_case_body(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        document_id: str,
    ) -> None:
        """Both field names are single words, so there is no camelCase variant."""
        response = client_with_s3.put(
            f"/guidance/documents/{document_id}/sections/1",
            json={"heading": "Overview", "markdown": "Text."},
        )

        assert response.status_code == 204

    def test_preserves_a_heading_that_begins_with_its_section_number(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
        document_id: str,
    ) -> None:
        """ "7 day rule" in section 7 is legitimate and must not be de-duplicated."""
        client_with_s3.put(
            f"/guidance/documents/{document_id}/sections/7",
            json={"heading": "7 day rule", "markdown": "Text."},
        )

        assert (
            mock_s3_repo.upload_section.await_args.args[2]
            == "## 7 7 day rule\n\nText.\n"
        )

    def test_returns_404_for_unknown_section_number(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
        document_id: str,
    ) -> None:
        response = client_with_s3.put(
            f"/guidance/documents/{document_id}/sections/9",
            json={"heading": "Nope", "markdown": "Text."},
        )

        assert response.status_code == 404
        mock_s3_repo.upload_section.assert_not_awaited()

    def test_returns_404_for_unknown_document(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        mock_s3_repo: AsyncMock,
        document_id: str,
    ) -> None:
        mock_s3_repo.download_manifest.side_effect = _no_such_key_error()

        response = client_with_s3.put(
            f"/guidance/documents/{document_id}/sections/1",
            json={"heading": "Overview", "markdown": "Text."},
        )

        assert response.status_code == 404

    def test_returns_422_for_invalid_section_number(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        document_id: str,
    ) -> None:
        response = client_with_s3.put(
            f"/guidance/documents/{document_id}/sections/not-a-number",
            json={"heading": "Overview", "markdown": "Text."},
        )

        assert response.status_code == 422

    def test_returns_422_for_invalid_document_id(
        self,
        client_with_s3: fastapi.testclient.TestClient,
    ) -> None:
        response = client_with_s3.put(
            "/guidance/documents/not-a-uuid/sections/1",
            json={"heading": "Overview", "markdown": "Text."},
        )

        assert response.status_code == 422

    def test_returns_422_for_empty_heading(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        document_id: str,
    ) -> None:
        response = client_with_s3.put(
            f"/guidance/documents/{document_id}/sections/1",
            json={"heading": "", "markdown": "Text."},
        )

        assert response.status_code == 422

    def test_returns_422_when_markdown_is_missing(
        self,
        client_with_s3: fastapi.testclient.TestClient,
        document_id: str,
    ) -> None:
        """A partial update would silently blank the body."""
        response = client_with_s3.put(
            f"/guidance/documents/{document_id}/sections/1",
            json={"heading": "Overview"},
        )

        assert response.status_code == 422


class TestSwaggerDocumentation:
    """Test that API documentation is available."""

    def test_swagger_docs_available(
        self, client_with_mocks: fastapi.testclient.TestClient
    ) -> None:
        """Test that Swagger docs are accessible at /docs."""
        response = client_with_mocks.get("/docs")
        assert response.status_code == 200

    def test_openapi_schema_includes_guidance_endpoints(
        self, client_with_mocks: fastapi.testclient.TestClient
    ) -> None:
        """Test that OpenAPI schema includes guidance endpoints."""
        response = client_with_mocks.get("/openapi.json")
        assert response.status_code == 200
        data = response.json()
        assert "/guidance/documents/" in data["paths"]
