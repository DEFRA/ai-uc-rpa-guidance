"""Tests for PipelineDocumentParser."""

import io
import struct
import zlib
from unittest.mock import AsyncMock

import bson
import docx
import pytest

from app.guidance.documents import models, parser, s3_repository


def _make_png_bytes() -> bytes:
    """Generate a valid minimal 1×1 white RGB PNG."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
    raw_row = b"\x00\xff\xff\xff"  # filter=None, R=255, G=255, B=255
    idat = chunk(b"IDAT", zlib.compress(raw_row))
    iend = chunk(b"IEND", b"")
    return sig + ihdr + idat + iend


def _make_minimal_docx(title: str = "") -> bytes:
    """Build a minimal in-memory .docx with a heading and a paragraph."""
    doc = docx.Document()
    if title:
        doc.core_properties.title = title
    doc.add_heading("Test Section", level=1)
    doc.add_paragraph("Hello world.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_images(count: int = 1) -> bytes:
    """Build a minimal .docx with headings and embedded PNG images."""
    doc = docx.Document()
    png = _make_png_bytes()
    for i in range(count):
        doc.add_heading(f"Section {i + 1}", level=1)
        doc.add_picture(io.BytesIO(png))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_inline_icon() -> bytes:
    """Build a .docx whose bullet embeds an icon mid-sentence: text, image, text."""
    doc = docx.Document()
    doc.add_heading("Steps", level=1)
    bullet = doc.add_paragraph(style="List Bullet")
    bullet.add_run("Select the ")
    bullet.add_run().add_picture(io.BytesIO(_make_png_bytes()))
    bullet.add_run("binocular icon")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_document(
    path: str = "s3://source-bucket/doc-id/file-id",
) -> models.GuidanceDocument:
    return models.GuidanceDocument(
        id=str(bson.ObjectId()),
        title="Test Doc",
        status=models.ExtractionStatus.PROCESSING,
        path=path,
    )


@pytest.fixture
def s3_repo() -> AsyncMock:
    repo = AsyncMock(spec=s3_repository.GuidanceS3Repository)
    repo.download_docx.return_value = _make_minimal_docx()
    return repo


@pytest.fixture
def parser_instance(s3_repo: AsyncMock) -> parser.PipelineDocumentParser:
    return parser.PipelineDocumentParser(s3_repo)


class TestPipelineDocumentParserSuccess:
    @pytest.mark.asyncio
    async def test_returns_complete_status(
        self,
        parser_instance: parser.PipelineDocumentParser,
    ) -> None:
        result = await parser_instance.parse(_make_document())

        assert result.status == models.ExtractionStatus.COMPLETE

    @pytest.mark.asyncio
    async def test_content_populated(
        self,
        parser_instance: parser.PipelineDocumentParser,
    ) -> None:
        result = await parser_instance.parse(_make_document())

        assert result.content is not None
        assert len(result.content) > 0
        assert "Test Section" in result.content

    @pytest.mark.asyncio
    async def test_error_message_none_on_success(
        self,
        parser_instance: parser.PipelineDocumentParser,
    ) -> None:
        result = await parser_instance.parse(_make_document())

        assert result.error_message is None

    @pytest.mark.asyncio
    async def test_markdown_uploaded_to_s3(
        self,
        parser_instance: parser.PipelineDocumentParser,
        s3_repo: AsyncMock,
    ) -> None:
        document = _make_document()
        result = await parser_instance.parse(document)

        s3_repo.upload_content.assert_called_once_with(document.id, result.content)

    @pytest.mark.asyncio
    async def test_manifest_uploaded_to_s3(
        self,
        parser_instance: parser.PipelineDocumentParser,
        s3_repo: AsyncMock,
    ) -> None:
        document = _make_document()
        await parser_instance.parse(document)

        s3_repo.upload_manifest.assert_called_once()
        call_args = s3_repo.upload_manifest.call_args
        assert call_args.args[0] == document.id
        import json

        manifest_data = json.loads(call_args.args[1])
        assert manifest_data["document_id"] == str(document.id)
        assert "sections" in manifest_data

    @pytest.mark.asyncio
    async def test_sections_uploaded_to_s3(
        self,
        parser_instance: parser.PipelineDocumentParser,
        s3_repo: AsyncMock,
    ) -> None:
        document = _make_document()
        await parser_instance.parse(document)

        assert s3_repo.upload_section.call_count >= 1
        first_call = s3_repo.upload_section.call_args_list[0]
        assert first_call.args[0] == document.id

    @pytest.mark.asyncio
    async def test_docx_downloaded_with_correct_key(
        self,
        parser_instance: parser.PipelineDocumentParser,
        s3_repo: AsyncMock,
    ) -> None:
        await parser_instance.parse(
            _make_document(path="s3://source-bucket/doc-id/file-id")
        )

        s3_repo.download_docx.assert_called_once_with("doc-id/file-id")

    @pytest.mark.asyncio
    async def test_title_inferred_from_docx_when_not_set(
        self,
        s3_repo: AsyncMock,
    ) -> None:
        s3_repo.download_docx.return_value = _make_minimal_docx(title="Inferred Title")
        parser_inst = parser.PipelineDocumentParser(s3_repo)
        document = models.GuidanceDocument(
            id=str(bson.ObjectId()),
            title=None,
            status=models.ExtractionStatus.PROCESSING,
            path="s3://source-bucket/doc-id/file-id",
        )
        result = await parser_inst.parse(document)

        assert result.title == "Inferred Title"

    @pytest.mark.asyncio
    async def test_title_none_when_already_set_on_document(
        self,
        s3_repo: AsyncMock,
    ) -> None:
        s3_repo.download_docx.return_value = _make_minimal_docx(title="Inferred Title")
        parser_inst = parser.PipelineDocumentParser(s3_repo)
        document = models.GuidanceDocument(
            id=str(bson.ObjectId()),
            title="Explicit Title",
            status=models.ExtractionStatus.PROCESSING,
            path="s3://source-bucket/doc-id/file-id",
        )
        result = await parser_inst.parse(document)

        assert result.title is None


class TestPipelineDocumentParserFailure:
    @pytest.mark.asyncio
    async def test_s3_error_returns_failed_status(
        self,
        parser_instance: parser.PipelineDocumentParser,
        s3_repo: AsyncMock,
    ) -> None:
        s3_repo.download_docx.side_effect = Exception("S3 connection error")

        result = await parser_instance.parse(_make_document())

        assert result.status == models.ExtractionStatus.FAILED
        assert "S3 connection error" in (result.error_message or "")

    @pytest.mark.asyncio
    async def test_missing_path_returns_failed_status(
        self,
        parser_instance: parser.PipelineDocumentParser,
    ) -> None:
        document = _make_document(path="")
        document.path = None

        result = await parser_instance.parse(document)

        assert result.status == models.ExtractionStatus.FAILED
        assert "has no path set" in (result.error_message or "")

    @pytest.mark.asyncio
    async def test_upload_error_returns_failed_status(
        self,
        parser_instance: parser.PipelineDocumentParser,
        s3_repo: AsyncMock,
    ) -> None:
        s3_repo.upload_content.side_effect = Exception("Upload failed")

        result = await parser_instance.parse(_make_document())

        assert result.status == models.ExtractionStatus.FAILED
        assert "Upload failed" in (result.error_message or "")

    @pytest.mark.asyncio
    async def test_failure_does_not_raise(
        self,
        parser_instance: parser.PipelineDocumentParser,
        s3_repo: AsyncMock,
    ) -> None:
        s3_repo.download_docx.side_effect = Exception("boom")

        result = await parser_instance.parse(_make_document())

        assert isinstance(result, parser.ParseResult)


class TestPipelineDocumentParserImages:
    @pytest.mark.asyncio
    async def test_no_images_no_upload_calls(
        self,
        parser_instance: parser.PipelineDocumentParser,
        s3_repo: AsyncMock,
    ) -> None:
        result = await parser_instance.parse(_make_document())

        assert result.status == models.ExtractionStatus.COMPLETE
        s3_repo.upload_image.assert_not_called()

    @pytest.mark.asyncio
    async def test_image_uploaded_to_s3(
        self,
        s3_repo: AsyncMock,
    ) -> None:
        s3_repo.download_docx.return_value = _make_docx_with_images(count=1)
        parser_inst = parser.PipelineDocumentParser(s3_repo)
        document = _make_document()

        result = await parser_inst.parse(document)

        assert result.status == models.ExtractionStatus.COMPLETE
        s3_repo.upload_image.assert_called_once()
        call_args = s3_repo.upload_image.call_args
        assert call_args.args[0] == document.id
        assert call_args.args[1] == "1_img_1.png"
        assert isinstance(call_args.args[2], bytes)
        assert call_args.args[3] == "image/png"

    @pytest.mark.asyncio
    async def test_image_rel_path_set_in_markdown(
        self,
        s3_repo: AsyncMock,
    ) -> None:
        s3_repo.download_docx.return_value = _make_docx_with_images(count=1)
        parser_inst = parser.PipelineDocumentParser(s3_repo)
        document = _make_document()

        result = await parser_inst.parse(document)

        assert result.content is not None
        assert f"/guidance/documents/{document.id}/images/1_img_1.png" in result.content

    @pytest.mark.asyncio
    async def test_multiple_images_all_uploaded(
        self,
        s3_repo: AsyncMock,
    ) -> None:
        s3_repo.download_docx.return_value = _make_docx_with_images(count=2)
        parser_inst = parser.PipelineDocumentParser(s3_repo)
        document = _make_document()

        await parser_inst.parse(document)

        assert s3_repo.upload_image.call_count == 2
        filenames = [c.args[1] for c in s3_repo.upload_image.call_args_list]
        assert filenames == ["1_img_1.png", "2_img_1.png"]

    @pytest.mark.asyncio
    async def test_inline_icon_uploaded_and_linked_in_markdown(
        self,
        s3_repo: AsyncMock,
    ) -> None:
        """An icon embedded mid-sentence is uploaded and rendered inline in the bullet."""
        s3_repo.download_docx.return_value = _make_docx_with_inline_icon()
        parser_inst = parser.PipelineDocumentParser(s3_repo)
        document = _make_document()

        result = await parser_inst.parse(document)

        assert result.status == models.ExtractionStatus.COMPLETE
        s3_repo.upload_image.assert_called_once()
        assert result.content is not None
        rel_path = f"/guidance/documents/{document.id}/images/1_img_1.png"
        assert f"Select the ![]({rel_path})binocular icon" in result.content

    @pytest.mark.asyncio
    async def test_image_upload_failure_returns_failed_status(
        self,
        s3_repo: AsyncMock,
    ) -> None:
        s3_repo.download_docx.return_value = _make_docx_with_images(count=1)
        s3_repo.upload_image.side_effect = Exception("S3 upload failed")
        parser_inst = parser.PipelineDocumentParser(s3_repo)

        result = await parser_inst.parse(_make_document())

        assert result.status == models.ExtractionStatus.FAILED
        assert "S3 upload failed" in (result.error_message or "")
