import struct
import zlib

from docx import Document
from docx.shared import Inches

from app.guidance.pipeline import models, service


def _make_png() -> bytes:
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr_crc = zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF
    ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", ihdr_crc)
    raw = b"\x00\xff\x00\x00"
    compressed = zlib.compress(raw)
    idat_crc = zlib.crc32(b"IDAT" + compressed) & 0xFFFFFFFF
    idat = (
        struct.pack(">I", len(compressed))
        + b"IDAT"
        + compressed
        + struct.pack(">I", idat_crc)
    )
    iend_crc = zlib.crc32(b"IEND") & 0xFFFFFFFF
    iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", iend_crc)
    return sig + ihdr + idat + iend


def _make_test_docx() -> Document:
    """Create a synthetic docx with headings, paragraphs, a table, and an image."""
    doc = Document()

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the intro paragraph.")

    doc.add_heading("Details", level=2)
    doc.add_paragraph("Some details here.")

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "100"
    table.cell(2, 0).text = "Beta"
    table.cell(2, 1).text = "200"

    doc.add_heading("More Info", level=2)
    doc.add_paragraph("Additional information.")

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("Final thoughts.")

    return doc


class TestParserStructure:
    def test_nested_sections(self):
        doc = _make_test_docx()
        tree = service.parse_doc(doc, title="TestDoc")

        assert tree.title == "TestDoc"
        assert len(tree.children) == 2
        assert tree.children[0].heading == "Introduction"
        assert tree.children[1].heading == "Conclusion"

    def test_subsections_nested(self):
        doc = _make_test_docx()
        tree = service.parse_doc(doc, title="TestDoc")

        intro = tree.children[0]
        assert len(intro.children) == 2
        assert intro.children[0].heading == "Details"
        assert intro.children[1].heading == "More Info"

    def test_section_numbers(self):
        doc = _make_test_docx()
        tree = service.parse_doc(doc, title="TestDoc")

        assert tree.children[0].number == "1"
        assert tree.children[0].children[0].number == "1.1"
        assert tree.children[0].children[1].number == "1.2"
        assert tree.children[1].number == "2"

    def test_paragraph_content(self):
        doc = _make_test_docx()
        tree = service.parse_doc(doc, title="TestDoc")

        intro = tree.children[0]
        paras = [n for n in intro.content if isinstance(n, models.ParagraphNode)]
        assert any(
            "This is the intro paragraph." in "".join(s.text for s in p.spans)
            for p in paras
        )

    def test_table_extraction(self):
        doc = _make_test_docx()
        tree = service.parse_doc(doc, title="TestDoc")

        details = tree.children[0].children[0]  # "Details" subsection
        tables = [n for n in details.content if isinstance(n, models.TableNode)]
        assert len(tables) == 1
        assert tables[0].headers == ["Name", "Value"]
        assert tables[0].rows == [["Alpha", "100"], ["Beta", "200"]]

    def test_single_cell_table_parsed_as_header_only(self):
        doc = Document()
        doc.add_heading("Section", level=1)
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Important note."
        tree = service.parse_doc(doc, title="T")
        tables = [
            n for n in tree.children[0].content if isinstance(n, models.TableNode)
        ]
        assert len(tables) == 1
        assert tables[0].headers == ["Important note."]
        assert tables[0].rows == []

    def test_fully_merged_table_parsed_as_header_only(self):
        doc = Document()
        doc.add_heading("Section", level=1)
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Callout content."
        table.cell(0, 0).merge(table.cell(1, 1))
        tree = service.parse_doc(doc, title="T")
        tables = [
            n for n in tree.children[0].content if isinstance(n, models.TableNode)
        ]
        assert len(tables) == 1
        assert tables[0].headers == ["Callout content."]
        assert tables[0].rows == []


class TestParserImages:
    def test_image_extraction(self, tmp_path):
        """Test that an inline image is extracted and held in memory."""
        doc = Document()
        doc.add_heading("With Image", level=1)

        png_path = tmp_path / "test.png"
        png_path.write_bytes(_make_png())
        doc.add_picture(str(png_path), width=Inches(1))

        doc.add_paragraph("After the image.")

        tree = service.parse_doc(doc, title="ImageTest")

        section = tree.children[0]
        images = [n for n in section.content if isinstance(n, models.ImageNode)]
        assert len(images) == 1
        # rel_path is blank until the caller uploads to S3
        assert images[0].rel_path == ""
        # Raw bytes are in memory
        assert len(images[0].data) > 0
        assert images[0].ext == ".png"

    def test_inline_image_preserved_in_list_item(self, tmp_path):
        """An icon embedded mid-sentence stays inline; its bullet and text survive.

        Mirrors the "Select the <icon> binocular icon" bullet in the SITI Tenure
        guidance: one bullet whose runs are text -> inline picture -> text. The
        image must not swallow the paragraph — the bullet stays a list item, and
        the words remain real text spans rather than being lost into an image's
        alt text.
        """
        png_path = tmp_path / "icon.png"
        png_path.write_bytes(_make_png())

        doc = Document()
        doc.add_heading("Steps", level=1)
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.add_run("Select the ")
        bullet.add_run().add_picture(str(png_path), width=Inches(0.2))
        bullet.add_run("binocular icon")

        tree = service.parse_doc(doc, title="InlineIcon")
        section = tree.children[0]

        # The bullet is a list item, not a block image promoted out of the list.
        lists = [n for n in section.content if isinstance(n, models.ListNode)]
        assert len(lists) == 1
        assert len(lists[0].items) == 1
        item = lists[0].items[0]

        # Spans are text -> inline image -> text, in document order.
        assert len(item.spans) == 3
        assert isinstance(item.spans[0], models.InlineSpan)
        assert isinstance(item.spans[1], models.ImageSpan)
        assert isinstance(item.spans[2], models.InlineSpan)
        assert item.spans[0].text == "Select the "
        assert item.spans[2].text == "binocular icon"
        assert len(item.spans[1].data) > 0

        # No block image was emitted for this bullet.
        assert not [n for n in section.content if isinstance(n, models.ImageNode)]

        # The words survive intact as text — not doubled, not lost to alt text.
        text = "".join(s.text for s in item.spans if isinstance(s, models.InlineSpan))
        assert text == "Select the binocular icon"


class TestParserLists:
    def test_bullet_list_extraction(self):
        """Test that bullet list paragraphs are grouped into a ListNode."""
        doc = Document()
        doc.add_heading("Steps", level=1)
        doc.add_paragraph("Step one", style="List Bullet")
        doc.add_paragraph("Step two", style="List Bullet")
        doc.add_paragraph("Sub step", style="List Bullet 2")
        doc.add_paragraph("Normal text after list.")

        tree = service.parse_doc(doc, title="ListTest")

        section = tree.children[0]
        lists = [n for n in section.content if isinstance(n, models.ListNode)]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]

        assert len(lists) == 1
        assert len(lists[0].items) == 3
        assert lists[0].items[0].level == 0
        assert lists[0].items[2].level == 1  # List Bullet 2 = nested
        assert len(paras) == 1


class TestParserInlineFormatting:
    def test_bold_spans(self):
        """Test that bold runs are captured in spans."""
        doc = Document()
        doc.add_heading("Formatted", level=1)
        para = doc.add_paragraph()
        para.add_run("Normal text ")
        para.add_run("bold text").bold = True
        para.add_run(" more normal")

        tree = service.parse_doc(doc, title="BoldTest")

        section = tree.children[0]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]
        assert len(paras) == 1
        spans = paras[0].spans
        bold_spans = [s for s in spans if s.bold]
        assert any("bold text" in s.text for s in bold_spans)

    def test_color_spans(self):
        """Test that font colour is captured in spans."""
        from docx.shared import RGBColor

        doc = Document()
        doc.add_heading("Coloured", level=1)
        para = doc.add_paragraph()
        run = para.add_run("red warning")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        tree = service.parse_doc(doc, title="ColorTest")

        section = tree.children[0]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]
        assert len(paras) == 1
        colored = [s for s in paras[0].spans if s.color]
        assert len(colored) == 1
        assert colored[0].color == "FF0000"

    def test_underline_spans(self):
        """Test that underlined runs are captured in spans."""
        doc = Document()
        doc.add_heading("Underlined", level=1)
        para = doc.add_paragraph()
        run = para.add_run("underlined text")
        run.underline = True

        tree = service.parse_doc(doc, title="UnderlineTest")

        section = tree.children[0]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]
        assert len(paras) == 1
        underlined = [s for s in paras[0].spans if s.underline]
        assert any("underlined text" in s.text for s in underlined)

    def test_bold_val_false_not_bold(self):
        """w:b w:val='false' must not be treated as bold (explicit-off overrides presence)."""
        from docx.oxml.ns import qn
        from lxml import etree

        doc = Document()
        doc.add_heading("Toggle", level=1)
        para = doc.add_paragraph()
        run = para.add_run("not bold")

        # Inject w:b w:val="false" directly into the run's rPr to simulate a
        # style that explicitly disables bold on an otherwise-bold base style.
        rpr = run._r.get_or_add_rPr()
        b_elem = etree.SubElement(rpr, qn("w:b"))
        b_elem.set(qn("w:val"), "false")

        tree = service.parse_doc(doc, title="ValFalseTest")

        section = tree.children[0]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]
        assert len(paras) == 1
        spans = paras[0].spans
        bold_spans = [s for s in spans if s.bold]
        assert not any("not bold" in s.text for s in bold_spans)


class TestTitleInference:
    def test_explicit_title_wins(self):
        doc = Document()
        doc.core_properties.title = "Core Title"
        doc.add_heading("Heading One", level=1)
        tree = service.parse_doc(doc, title="Explicit Title")
        assert tree.title == "Explicit Title"

    def test_infers_from_core_properties(self):
        doc = Document()
        doc.core_properties.title = "Core Title"
        doc.add_heading("Heading One", level=1)
        tree = service.parse_doc(doc)
        assert tree.title == "Core Title"

    def test_infers_from_title_style(self):
        doc = Document()
        doc.core_properties.title = ""
        para = doc.add_paragraph("Style Title")
        para.style = doc.styles["Title"]
        doc.add_heading("Heading One", level=1)
        tree = service.parse_doc(doc)
        assert tree.title == "Style Title"

    def test_infers_from_first_page_plain_text(self):
        doc = Document()
        doc.core_properties.title = ""
        doc.add_paragraph("Plain Title Text")
        doc.add_paragraph("Plain Subtitle Text")
        tree = service.parse_doc(doc)
        assert tree.title == "Plain Title Text — Plain Subtitle Text"

    def test_infers_from_first_page_mixed_styles(self):
        doc = Document()
        doc.core_properties.title = ""
        doc.add_heading("Main Title", level=1)
        doc.add_paragraph("Plain subtitle")
        tree = service.parse_doc(doc)
        assert tree.title == "Main Title — Plain subtitle"

    def test_content_after_page_break_excluded(self):
        from docx.oxml.ns import qn
        from lxml import etree

        doc = Document()
        doc.core_properties.title = ""
        doc.add_paragraph("Cover Title")

        # Insert a manual page break
        break_para = doc.add_paragraph()
        run = break_para.add_run()
        br = etree.SubElement(run._r, qn("w:br"))
        br.set(qn("w:type"), "page")

        doc.add_paragraph("Body content not part of title")
        tree = service.parse_doc(doc)
        assert tree.title == "Cover Title"

    def test_empty_title_when_no_content(self):
        doc = Document()
        doc.core_properties.title = ""
        tree = service.parse_doc(doc)
        assert tree.title == ""

    def test_extract_title_core_properties(self):
        doc = Document()
        doc.core_properties.title = "  Trimmed Title  "
        assert service.DocxParser._extract_title(doc) == "Trimmed Title"

    def test_template_title_falls_back_to_heading1(self):
        doc = Document()
        doc.core_properties.title = "Design Team Guidance Template"
        doc.add_heading("Real Document Title", level=1)
        tree = service.parse_doc(doc)
        assert tree.title == "Real Document Title"

    def test_template_title_falls_back_to_title_style(self):
        doc = Document()
        doc.core_properties.title = "Design Team Guidance Template"
        para = doc.add_paragraph("Style Title")
        para.style = doc.styles["Title"]
        doc.add_heading("Heading One", level=1)
        tree = service.parse_doc(doc)
        assert tree.title == "Style Title"


def _add_bookmark(paragraph, name: str, bookmark_id: int = 0) -> None:
    """Wrap a bookmarkStart/End with the given name into a paragraph element.

    Mirrors how Word marks a cross-reference target (e.g. on a heading).
    """
    from docx.oxml.ns import qn
    from lxml import etree

    start = etree.SubElement(paragraph._element, qn("w:bookmarkStart"))
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = etree.SubElement(paragraph._element, qn("w:bookmarkEnd"))
    end.set(qn("w:id"), str(bookmark_id))


def _add_anchor_link(paragraph, anchor: str, text: str) -> None:
    """Append a w:hyperlink with only a w:anchor (an internal cross-reference)."""
    from docx.oxml.ns import qn
    from lxml import etree

    hyperlink = etree.SubElement(paragraph._element, qn("w:hyperlink"))
    hyperlink.set(qn("w:anchor"), anchor)
    run = etree.SubElement(hyperlink, qn("w:r"))
    t_elem = etree.SubElement(run, qn("w:t"))
    t_elem.text = text


class TestParserHyperlinks:
    def test_hyperlink_in_span(self):
        """w:hyperlink elements in a paragraph produce InlineSpan.hyperlink with the target URL."""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml.ns import qn
        from lxml import etree

        url = "https://example.com"

        doc = Document()
        doc.add_heading("Links", level=1)
        para = doc.add_paragraph()

        # Register an external hyperlink relationship on the paragraph's part.
        r_id = para.part.rels.get_or_add_ext_rel(RT.HYPERLINK, url)

        # Build a w:hyperlink element with the relationship id and a run inside it.
        hyperlink_elem = etree.SubElement(para._element, qn("w:hyperlink"))
        hyperlink_elem.set(qn("r:id"), r_id)
        run_elem = etree.SubElement(hyperlink_elem, qn("w:r"))
        t_elem = etree.SubElement(run_elem, qn("w:t"))
        t_elem.text = "click here"

        tree = service.parse_doc(doc, title="HyperlinkTest")

        section = tree.children[0]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]
        assert len(paras) == 1
        hyperlinked = [s for s in paras[0].spans if s.hyperlink]
        assert len(hyperlinked) == 1
        assert hyperlinked[0].text == "click here"
        assert hyperlinked[0].hyperlink == url

    def test_split_hyperlink_runs_coalesce_to_one_link(self):
        """A single w:hyperlink whose anchor text Word split across several runs must
        yield ONE span / ONE Markdown link, not one per run.

        Mirrors the real source XML in sections 4.1/4.2: a lone <w:hyperlink> element
        containing three <w:r> runs ("SFI", " ", "SITI Agri Basic Navigation"), each
        with an rStyle=Hyperlink rPr. Parsing per-run previously emitted three adjacent
        links to the identical URL.
        """
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml.ns import qn
        from lxml import etree

        from app.guidance.pipeline.renderers import markdown as markdown_renderer

        url = "https://example.com/siti-agri-basic-navigation-guide"

        doc = Document()
        doc.add_heading("Links", level=1)
        para = doc.add_paragraph()

        r_id = para.part.rels.get_or_add_ext_rel(RT.HYPERLINK, url)
        hyperlink_elem = etree.SubElement(para._element, qn("w:hyperlink"))
        hyperlink_elem.set(qn("r:id"), r_id)
        for fragment in ("SFI", " ", "SITI Agri Basic Navigation"):
            run_elem = etree.SubElement(hyperlink_elem, qn("w:r"))
            rpr = etree.SubElement(run_elem, qn("w:rPr"))
            style = etree.SubElement(rpr, qn("w:rStyle"))
            style.set(qn("w:val"), "Hyperlink")
            t_elem = etree.SubElement(run_elem, qn("w:t"))
            t_elem.set(qn("xml:space"), "preserve")
            t_elem.text = fragment

        tree = service.parse_doc(doc, title="SplitHyperlinkTest")

        section = tree.children[0]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]
        assert len(paras) == 1
        hyperlinked = [s for s in paras[0].spans if s.hyperlink]
        assert len(hyperlinked) == 1
        assert hyperlinked[0].text == "SFI SITI Agri Basic Navigation"
        assert hyperlinked[0].hyperlink == url

        # The user-visible goal: exactly one Markdown link, not three.
        rendered = markdown_renderer.section_to_markdown(section)
        assert rendered.count(f"](<{url}>)") == 1

    def test_empty_hyperlink_produces_no_span(self):
        """A w:hyperlink carrying no text (its runs have no w:t) yields no span, and
        does not disturb the surrounding runs."""
        from docx.oxml.ns import qn
        from lxml import etree

        doc = Document()
        doc.add_heading("Links", level=1)
        para = doc.add_paragraph()
        para.add_run("visible text")

        # A hyperlink element with a run but no w:t content — nothing to render.
        hyperlink_elem = etree.SubElement(para._element, qn("w:hyperlink"))
        hyperlink_elem.set(qn("w:anchor"), "_Somewhere")
        etree.SubElement(hyperlink_elem, qn("w:r"))

        tree = service.parse_doc(doc, title="EmptyHyperlinkTest")

        section = tree.children[0]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]
        assert len(paras) == 1
        spans = paras[0].spans
        assert [s.text for s in spans] == ["visible text"]
        assert all(s.hyperlink == "" for s in spans)

    def test_empty_bare_run_produces_no_span(self):
        """A bare run with no w:t text contributes no span; real runs are unaffected."""
        from docx.oxml.ns import qn
        from lxml import etree

        doc = Document()
        doc.add_heading("Empty", level=1)
        para = doc.add_paragraph()
        etree.SubElement(para._element, qn("w:r"))  # empty run, no w:t
        para.add_run("real text")

        tree = service.parse_doc(doc, title="EmptyRunTest")

        section = tree.children[0]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]
        assert len(paras) == 1
        assert [s.text for s in paras[0].spans] == ["real text"]

    def test_hyperlink_survives_adjacent_plain_runs(self):
        """A hyperlink-wrapped run sandwiched between plain runs must not leak onto
        its neighbours (regression for the id()-collision bug: the hyperlink map was
        keyed by id(run_elem), and lxml recycles proxy object ids across separate
        element traversals, silently misattributing links to unrelated runs)."""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml.ns import qn
        from lxml import etree

        url = "https://example.com/guide"

        doc = Document()
        doc.add_heading("Links", level=1)
        para = doc.add_paragraph()

        para.add_run("Before text ")

        r_id = para.part.rels.get_or_add_ext_rel(RT.HYPERLINK, url)
        hyperlink_elem = etree.SubElement(para._element, qn("w:hyperlink"))
        hyperlink_elem.set(qn("r:id"), r_id)
        run_elem = etree.SubElement(hyperlink_elem, qn("w:r"))
        t_elem = etree.SubElement(run_elem, qn("w:t"))
        t_elem.text = "the linked text"

        para.add_run(" after text")

        tree = service.parse_doc(doc, title="AdjacentHyperlinkTest")

        section = tree.children[0]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]
        assert len(paras) == 1
        spans = paras[0].spans
        assert len(spans) == 3

        before, linked, after = spans
        assert before.text == "Before text "
        assert before.hyperlink == ""
        assert linked.text == "the linked text"
        assert linked.hyperlink == url
        assert after.text == " after text"
        assert after.hyperlink == ""

    def test_internal_link_resolves_to_section_number(self):
        """A cross-reference to a bookmark on a later heading is rewritten from the
        raw bookmark anchor to that section's positional number."""
        doc = Document()
        doc.add_heading("Intro", level=1)  # section 1
        para = doc.add_paragraph()
        _add_anchor_link(para, "_Foo_1", "go to foo")

        foo = doc.add_heading("Foo bar", level=1)  # section 2
        _add_bookmark(foo, "_Foo_1")
        doc.add_paragraph("Foo body.")

        tree = service.parse_doc(doc, title="InternalLinkTest")

        intro = tree.children[0]
        paras = [n for n in intro.content if isinstance(n, models.ParagraphNode)]
        hyperlinked = [s for s in paras[0].spans if s.hyperlink]
        assert len(hyperlinked) == 1
        assert hyperlinked[0].text == "go to foo"
        assert hyperlinked[0].hyperlink == "#2"

    def test_internal_link_resolves_renamed_bookmark(self):
        """Resolution is by bookmark identity, not heading text: a bookmark whose
        name bears no textual relation to its heading still resolves. This is the
        case a frontend slug heuristic can never recover."""
        doc = Document()
        doc.add_heading("First", level=1)  # section 1
        para = doc.add_paragraph()
        _add_anchor_link(para, "_Alias", "see the other section")

        target = doc.add_heading("Totally Different", level=1)  # section 2
        _add_bookmark(target, "_Alias")
        doc.add_paragraph("Target body.")

        tree = service.parse_doc(doc, title="RenamedBookmarkTest")

        first = tree.children[0]
        paras = [n for n in first.content if isinstance(n, models.ParagraphNode)]
        hyperlinked = [s for s in paras[0].spans if s.hyperlink]
        assert len(hyperlinked) == 1
        assert hyperlinked[0].hyperlink == "#2"

    def test_internal_link_unresolvable_anchor_preserved_and_logged(self, caplog):
        """An anchor with no matching bookmarkStart (e.g. a bare TOC anchor) is left
        untouched and logged — an unresolved cross-reference is a data signal, not
        something to drop silently."""
        import logging

        doc = Document()
        doc.add_heading("Links", level=1)
        para = doc.add_paragraph()
        _add_anchor_link(para, "_Toc123456", "back to top")

        with caplog.at_level(logging.WARNING):
            tree = service.parse_doc(doc, title="AnchorLinkTest")

        section = tree.children[0]
        paras = [n for n in section.content if isinstance(n, models.ParagraphNode)]
        assert len(paras) == 1
        hyperlinked = [s for s in paras[0].spans if s.hyperlink]
        assert len(hyperlinked) == 1
        assert hyperlinked[0].text == "back to top"
        assert hyperlinked[0].hyperlink == "#_Toc123456"
        assert "#_Toc123456" in caplog.text


class TestParserImageEdgeCases:
    def test_image_missing_embed_id(self, tmp_path):
        """A blip element with no r:embed attribute is silently skipped — no ImageNode, no error."""
        from docx.oxml.ns import qn

        doc = Document()
        doc.add_heading("Section", level=1)

        png_path = tmp_path / "test.png"
        png_path.write_bytes(_make_png())
        doc.add_picture(str(png_path), width=Inches(1))

        # Find the blip element and strip its r:embed attribute.
        body = doc.element.body
        blip = body.find(f".//{qn('a:blip')}")
        assert blip is not None
        embed_attr = qn("r:embed")
        if embed_attr in blip.attrib:
            del blip.attrib[embed_attr]

        tree = service.parse_doc(doc, title="NoEmbedTest")

        section = tree.children[0]
        images = [n for n in section.content if isinstance(n, models.ImageNode)]
        assert len(images) == 0

    def test_image_missing_relationship(self, tmp_path):
        """A blip with an r:embed that has no matching rel is silently skipped."""
        from docx.oxml.ns import qn

        doc = Document()
        doc.add_heading("Section", level=1)

        png_path = tmp_path / "test.png"
        png_path.write_bytes(_make_png())
        doc.add_picture(str(png_path), width=Inches(1))

        # Replace the r:embed value with an ID that doesn't exist in part.rels.
        body = doc.element.body
        blip = body.find(f".//{qn('a:blip')}")
        assert blip is not None
        blip.set(qn("r:embed"), "rId999")

        tree = service.parse_doc(doc, title="MissingRelTest")

        section = tree.children[0]
        images = [n for n in section.content if isinstance(n, models.ImageNode)]
        assert len(images) == 0


class TestParserContentOutsideSection:
    def test_table_before_heading_ignored(self):
        """A table that appears before any heading is silently discarded (empty stack)."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "H1"
        table.cell(0, 1).text = "H2"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "B"

        tree = service.parse_doc(doc, title="TableBeforeHeading")

        assert tree.children == []

    def test_image_before_heading_ignored(self, tmp_path):
        """An image that appears before any heading is silently discarded (empty stack)."""
        doc = Document()
        png_path = tmp_path / "test.png"
        png_path.write_bytes(_make_png())
        doc.add_picture(str(png_path), width=Inches(1))

        tree = service.parse_doc(doc, title="ImageBeforeHeading")

        assert tree.children == []

    def test_list_before_heading_ignored(self):
        """List items before the first heading are discarded when the stack is flushed empty."""
        doc = Document()
        doc.add_paragraph("Item one", style="List Bullet")
        doc.add_paragraph("Item two", style="List Bullet")
        # Adding a heading flushes the pending list with an empty stack.
        doc.add_heading("First Section", level=1)

        tree = service.parse_doc(doc, title="ListBeforeHeading")

        # The section exists but has no list content from the pre-heading items.
        assert len(tree.children) == 1
        section = tree.children[0]
        lists = [n for n in section.content if isinstance(n, models.ListNode)]
        assert lists == []


class TestParserListLevel:
    def test_numpr_ilvl_level(self):
        """A List Paragraph with w:numPr/w:ilvl w:val='2' reports level 2."""
        from docx.oxml.ns import qn
        from lxml import etree

        doc = Document()
        doc.add_heading("Numbered", level=1)
        para = doc.add_paragraph("Deep item", style="List Paragraph")

        # Inject w:numPr containing w:ilvl w:val="2" and a stub w:numId.
        ppr = para._element.get_or_add_pPr()
        num_pr = etree.SubElement(ppr, qn("w:numPr"))
        ilvl = etree.SubElement(num_pr, qn("w:ilvl"))
        ilvl.set(qn("w:val"), "2")
        num_id = etree.SubElement(num_pr, qn("w:numId"))
        num_id.set(qn("w:val"), "1")

        tree = service.parse_doc(doc, title="IlvlTest")

        section = tree.children[0]
        lists = [n for n in section.content if isinstance(n, models.ListNode)]
        assert len(lists) == 1
        assert lists[0].items[0].level == 2
